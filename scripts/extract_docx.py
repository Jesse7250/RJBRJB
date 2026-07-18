from docx import Document

path = r'D:\Gitproject\RJBRJB\基于大模型的个性化资源生成与学习多智能体系统开发.docx'
doc = Document(path)
out = []
for i, para in enumerate(doc.paragraphs, 1):
    text = para.text.strip()
    if text:
        out.append(f'{i:04d}: {text}')
for table_idx, table in enumerate(doc.tables, 1):
    out.append(f'\n[TABLE {table_idx}]')
    for row in table.rows:
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        out.append(' | '.join(cells))
text = '\n'.join(out)
with open(r'D:\Gitproject\RJBRJB\docs\比赛要求_extracted.txt', 'w', encoding='utf-8') as f:
    f.write(text)
print(f'Extracted {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables')
print(text[:1000])
